#!/usr/bin/env python3
"""
Analyzing Islam Vol I — Word Structural Prototype
B5 (176×250mm), mirrored margins, 262 deduplicated Quran entries.
Run: python build-book-docx.py
"""
import re, json, html as html_mod
from pathlib import Path
from docx import Document
from docx.shared import Mm, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE    = Path(r"C:\Users\zande\Documents\AI Workspace\Analyzing Islam")
CATALOG = BASE / "site/assets/data/catalog-entries.json"
QURAN   = BASE / "site/catalog/quran.html"
OUT_DIR = BASE / "book-design/vol1-quran"
OUT     = OUT_DIR / "Analyzing Islam Vol I — Word Prototype.docx"

# ── Exclusions ────────────────────────────────────────────────────────────────
EXCLUDE_IDS = {
    # Original duplicates
    "amputate-the-hand-of-the-thief-regardless-of-circumstance-4104d45b",
    "one-hundred-lashes-for-fornication-yet-the-hadith-demands-st-f805f912",
    # Hard duplicates — thinner entry removed in each pair
    "jews-transformed-into-apes-a205d9d7",
    "the-sun-runs-to-a-fixed-resting-place-5f69c2e2",
    "quran-fire-punishment-to-skin-replace",
    "polytheists-are-unclean-and-forbidden-from-the-sacred-mosque-793234d0",
    "fabricated-quote-jews-say-ezra-is-the-son-of-allah-df9200f3",
    "quran-menstruating-retreat",
    "quran-cow-that-killed",
    "quran-iblis-command-prostrate",
    "jinn-listen-to-the-quran-in-a-tree-and-convert-63828ff4",
    "creation-in-six-days-or-eight-a-day-count-contradiction-201b57cd",
    "quran-predestination-but-punishment",
    "quran-allah-best-plotters-jesus",
    "quran-pharaoh-wall-building",
    "quran-do-not-befriend-kafir",
    # Near-duplicates — user decisions
    "quran-right-hand-sex-captive-wife",
    "quran-children-spoils-war",
    "quran-number-of-sleepers",
    "quran-how-long-sleepers-slept",
    "quran-muhammad-mutah-private-wife",
    "quran-prophet-captives-war-booty",
}

# ── Chapter definitions ───────────────────────────────────────────────────────
CHAPTERS = {
    1:  ("Abrogation",
         "The doctrine of naskh (abrogation) holds that later Quranic verses can supersede earlier ones "
         "while both remain in the written text. The Quran references this principle explicitly at Q 2:106 and "
         "Q 16:101. This chapter examines the theological and logical problems that arise when a supposedly "
         "perfect, eternal divine text requires internal revision — and the implications for claims of "
         "divine omniscience and consistency."),
    2:  ("Scripture Integrity",
         "The Quran presents itself as a perfectly preserved, uniquely clear, and self-authenticating revelation. "
         "This chapter examines the passages and historical facts that challenge those claims: the burning of variant "
         "codices under Uthman, the acknowledgement of verses no one knows the interpretation of, the Islamic Dilemma "
         "between Quranic endorsement of the Bible and the doctrine of tahrif, and internal claims that "
         "undermine the text's own standard of clarity and permanence."),
    3:  ("Contradictions",
         "A text claimed to be the direct word of an omniscient God is expected to be internally consistent. "
         "This chapter catalogues passages in the Quran that contradict other Quranic passages — in "
         "arithmetic, cosmology, prophetic history, and law — and examines the apologetic strategies "
         "employed to reconcile them."),
    4:  ("Logical Inconsistency",
         "Beyond factual contradiction, several Quranic passages generate problems of logical form: self-refuting "
         "claims, arguments that assume what they are meant to prove, divine attributes that cannot coherently "
         "coexist, and instructions whose application requires information the text withholds."),
    5:  ("Allah's Character",
         "Islamic theology attributes to Allah a set of perfections — omniscience, omnipotence, justice, mercy. "
         "A number of Quranic passages sit in tension with one or more of those attributes: a God who seals "
         "disbelievers' hearts and then punishes them for disbelief, who engineers deception, who creates "
         "beings destined for Hell."),
    6:  ("Cosmology",
         "A number of Quranic passages describe the physical universe in ways that reflect pre-scientific "
         "cosmological assumptions rather than observed reality — including a sun that sets in a muddy "
         "spring, a seven-layered sky, sperm produced between the backbone and ribs, and creation in six days."),
    7:  ("Pre-Islamic Borrowings",
         "Several Quranic narratives have direct parallels in Jewish midrashic literature, Christian apocryphal "
         "gospels, Zoroastrian texts, and pre-Islamic Arabian legend — including stories, characters, "
         "and details not found in canonical Jewish or Christian scripture."),
    8:  ("Prophetic Character",
         "The Quran presents Muhammad as the exemplary moral model (uswa hasana). Several passages, however, "
         "describe a prophet who required divine reassurance, who benefited personally from his own revelations, "
         "who was rebuked by Allah for specific decisions, and whose conduct raises ethical questions the text "
         "itself registers without resolving."),
    9:  ("Prophetic Privileges",
         "A cluster of Quranic verses grants Muhammad exemptions and permissions explicitly denied to ordinary "
         "believers — additional wives beyond the limit of four, marriage to his adopted son's divorcee, "
         "a personal cut of war spoils, permission for women to offer themselves to him without a dowry."),
    10: ("Jesus / Christology",
         "The Quran contains a substantial Christology — an account of Jesus that agrees with some Christian "
         "claims, categorically denies others, and adds details found in no earlier canonical source."),
    11: ("Women & Sexual Issues",
         "The Quran legislates extensively on the status of women, marriage, sexual access, and related matters. "
         "Several passages establish legal and social hierarchies that contemporary ethics regards as "
         "discriminatory, permit practices modern law criminalises, or create internal inconsistencies."),
    12: ("Child Marriage",
         "Q 65:4 sets out divorce procedures for wives who have not yet menstruated — an explicit "
         "Quranic provision for marriage to pre-pubescent girls. This chapter examines the verse, its classical "
         "tafsir, the apologetic strategies used to contextualise it, and why those strategies do not remove "
         "the ethical problem the text creates."),
    13: ("LGBTQ / Gender",
         "The Quran's account of Lot's people and related passages have been read by the classical tradition "
         "as a divine condemnation of same-sex relations. This chapter examines those passages, the related "
         "verse on gender non-conformity, and the ethical problems they raise."),
    14: ("Slavery & Captives",
         "The Quran regulates slavery rather than prohibiting it — specifying procedures for manumission, "
         "permitting sexual access to female captives, and treating enslaved people as a recognised legal "
         "category."),
    15: ("Warfare & Jihad",
         "Several Quranic verses command violence against non-Muslims in terms that admit no obvious limiting "
         "context — commanding believers to kill, fight, or subjugate until conversion or submission is obtained."),
    16: ("Apostasy & Blasphemy",
         "The Quran does not state an explicit death penalty for apostasy, but several passages are read by "
         "classical jurists as endorsing it, and Q 4:89 is the key proof-text for that ruling."),
    17: ("Governance",
         "A number of passages establish that sovereignty belongs to Allah alone and that legislation is his "
         "exclusive prerogative — the canonical proof-texts for Islamic theocratic governance."),
    18: ("Disbelievers & Moral Problems",
         "The Quran characterises non-Muslims in terms ranging from misguided to irredeemably corrupt, the worst "
         "of creatures, and objects of divine curse. Several passages mandate social and legal discrimination "
         "against them."),
    19: ("Antisemitism",
         "The Quran contains direct derogatory characterisations of Jews as a group: divine transformation into "
         "apes and pigs, fabricated theological claims attributed to them, and placement at the top of a ranking "
         "of hostility toward believers."),
    20: ("Paradise",
         "The Quran's descriptions of paradise are detailed and physical: gardens of flowing rivers, the "
         "eternal virgin houris, rivers of wine and honey. This chapter examines passages where those "
         "descriptions raise moral problems."),
    21: ("Strange",
         "A number of Quranic passages describe supernatural events, historical claims, or cosmological "
         "assertions that resist straightforward naturalisation — stars as missiles thrown at eavesdropping "
         "jinn, villages left dead for a century, ants that converse with Solomon."),
    22: ("Magic & Ritual",
         "The Quran legislates extensively on ritual purity and acknowledges a world populated by jinn, "
         "sorcerers, and supernatural entities. Several passages describe magic as real and dangerous."),
    23: ("Animals",
         "Several Quranic passages about animals create scientific, moral, or theological problems: bees that "
         "receive divine inspiration, animals that form communities like humans, Solomon commanding the ants."),
}

# TAG_PRIORITY: 31 priority-ordered (tag, chapter_num) pairs
TAG_PRIORITY = [
    ("antisemitism", 19), ("childmarriage", 12), ("apostasy", 16),
    ("privileges",    9), ("jesus",         10), ("abrogation",   1),
    ("warfare",      15), ("governance",    17), ("preislamic",   7),
    ("scripture",     2), ("allah",          5), ("cosmology",    6),
    ("science",       6), ("magic",         22), ("slavery",     14),
    ("women",        11), ("sexual",        11), ("hudud",       17),
    ("prophet",       8), ("animals",       23), ("ritual",      22),
    ("contradiction", 3), ("logic",          4), ("paradise",    20),
    ("hell",         20), ("disbelievers",  18), ("morality",    18),
    ("strange",      21), ("incest",        13), ("gross-vile",  13),
    ("lgbtq",        13),
]

ID_OVERRIDES = {
    "the-seven-sleepers-of-ephesus-a-christian-legend-as-quranic-13829e66": 7,
    "sexual-access-to-married-female-slaves-right-hand-possesses-25cd8f4b": 11,
    "quran-wudu-tayammum-touching-women": 22,
    "quran-abasa-frowned-blind-rebuke": 8,
    "quran-inheritance-fractions-do-not-sum": 4,
    "quran-46-15-31-14-six-month-gestation-arithmetic": 6,
    "quran-69-32-seventy-cubit-chain-ghislin-food": 20,
    "paradise-as-physical-pleasure-garden-with-purified-spouses-65756a43": 20,
    "the-houris-eternal-virgins-as-paradise-reward-d8c254e9": 20,
    "quran-quran-as-healing": 2,
    "islamic-dilemma": 2,
    "the-quran-endorses-jews-and-christians-to-judge-by-their-own-32929162": 2,
    "no-one-can-change-the-words-of-allah-yet-tahrif-is-the-centr-d98f36e4": 2,
    # Chapter population fixes — ensure ch 14, 23 are non-empty
    # NOTE: zaynab-affair was incorrectly overridden to ch 13 (LGBTQ/Gender);
    # it belongs in ch 9 (Prophetic Privileges) via its 'privileges' tag.
    "prophet-should-not-take-captives-until-he-inflicts-a-massacr-75d23fb1": 14,
    "quran-38-31-33-solomon-hamstrings-the-horses": 23,
}

STRENGTH_ORDER = {"basic": 0, "moderate": 1, "strong": 2}

# ── Helpers ───────────────────────────────────────────────────────────────────

def strip_tags(s):
    """Strip HTML tags and decode common entities."""
    s = re.sub(r'<[^>]+>', '', s)
    for ent, ch in [('&amp;','&'),('&lt;','<'),('&gt;','>'),('&nbsp;',' '),
                    ('&#8212;','—'),('&#8211;','–'),('&#8216;',"'"),('&#8217;',"'"),
                    ('&#8220;','"'),('&#8221;','"'),('&mdash;','—'),('&ndash;','–'),
                    ('&rsquo;',"'"),('&lsquo;',"'"),('&ldquo;','"'),('&rdquo;','"'),
                    ('&hellip;','…')]:
        s = s.replace(ent, ch)
    return re.sub(r'[ \t]+', ' ', s).strip()

def assign_chapter(eid, categories):
    if eid in ID_OVERRIDES:
        return ID_OVERRIDES[eid]
    for tag, ch in TAG_PRIORITY:
        if tag in categories:
            return ch
    return 18

def parse_entries():
    """Parse full entry content from quran.html. Returns dict[id -> sections]."""
    raw = QURAN.read_text(encoding='utf-8', errors='ignore')
    pat = r'<div[^>]+class="[^"]*\bentry\b[^"]*"[^>]+id="([^"]+)"[^>]*>'
    opens = list(re.finditer(pat, raw))
    result = {}
    for i, m in enumerate(opens):
        eid = m.group(1)
        end = opens[i+1].start() if i+1 < len(opens) else len(raw)
        chunk = raw[m.start():end]
        bq_m = re.search(r'<blockquote[^>]*>(.*?)</blockquote>', chunk, re.DOTALL)
        quote = ''
        if bq_m:
            q = re.sub(r'<p[^>]*>', '', bq_m.group(1))
            q = re.sub(r'</p>', ' ', q)
            quote = strip_tags(q).strip()
        h4_parts = re.split(r'<h4[^>]*>', chunk)
        sections = {'quote': quote, 'says': '', 'problem': '', 'response': '', 'fails': ''}
        for part in h4_parts[1:]:
            end_tag = part.find('</h4>')
            if end_tag == -1:
                continue
            header = part[:end_tag].lower().strip()
            body = part[end_tag+5:]
            paras = re.findall(r'<p[^>]*>(.*?)</p>', body, re.DOTALL)
            text = '\n\n'.join(strip_tags(p) for p in paras if p.strip())
            if 'what the verse' in header:
                sections['says'] = text
            elif 'why this is a problem' in header:
                sections['problem'] = text
            elif 'muslim response' in header:
                sections['response'] = text
            elif 'why it fails' in header:
                sections['fails'] = text
        result[eid] = sections
    return result

def get_entries():
    """Return catalog entries filtered to active Quran entries only."""
    catalog = json.loads(CATALOG.read_text(encoding='utf-8'))
    return [e for e in catalog
            if e.get('source') == 'quran' and e['id'] not in EXCLUDE_IDS]

def build_chapters(entries):
    """Assign entries to chapters and sort basic->moderate->strong."""
    chapters = {n: [] for n in CHAPTERS}
    for e in entries:
        ch = assign_chapter(e['id'], e.get('categories', []))
        chapters[ch].append(e)
    for ch in chapters:
        chapters[ch].sort(key=lambda e: STRENGTH_ORDER.get(e.get('strength',''), 0))
    return chapters


def setup_document(doc):
    """Configure B5 page size, mirrored margins, and footer."""
    section = doc.sections[0]
    section.page_width  = Mm(176)
    section.page_height = Mm(250)
    section.top_margin    = Mm(20)
    section.bottom_margin = Mm(25)
    section.left_margin   = Mm(25)   # inner (gutter side)
    section.right_margin  = Mm(18)   # outer

    # Enable mirror margins via document settings XML
    settings = doc.settings.element
    mirror = OxmlElement('w:mirrorMargins')
    settings.append(mirror)

    setup_footer(section)


def setup_footer(section):
    """Add centred Arabic page number to footer."""
    section.footer_distance = Mm(12)
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()

    r1 = para.add_run()
    fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin')
    r1._r.append(fc1)

    r2 = para.add_run()
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = ' PAGE '
    r2._r.append(it)

    r3 = para.add_run()
    fc3 = OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'), 'separate')
    r3._r.append(fc3)

    r4 = para.add_run('1')   # placeholder shown before field update

    r5 = para.add_run()
    fc5 = OxmlElement('w:fldChar'); fc5.set(qn('w:fldCharType'), 'end')
    r5._r.append(fc5)


def setup_styles(doc):
    """Define all AI_* paragraph styles used in the document."""

    def _make(name, font_name, size_pt, bold=False, italic=False,
               before_pt=0, after_pt=6, line_spacing_pt=None,
               align=WD_ALIGN_PARAGRAPH.LEFT,
               keep_next=False, page_break_before=False,
               color_rgb=None):
        try:
            s = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            s = doc.styles[name]
        f = s.font
        f.name = font_name
        f.size = Pt(size_pt)
        f.bold = bold
        f.italic = italic
        if color_rgb:
            f.color.rgb = RGBColor(*color_rgb)
        pf = s.paragraph_format
        pf.space_before = Pt(before_pt)
        pf.space_after  = Pt(after_pt)
        pf.alignment    = align
        pf.keep_with_next = keep_next
        pf.page_break_before = page_break_before
        if line_spacing_pt:
            pf.line_spacing = Pt(line_spacing_pt)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        return s

    # Body text
    _make('AI_Normal',       'Georgia', 11, after_pt=6,  line_spacing_pt=14.5)
    # Entry elements
    _make('AI_Breadcrumb',   'Calibri',  8, after_pt=4,  color_rgb=(120,120,120))
    _make('AI_EntryTitle',   'Georgia', 14, bold=True, before_pt=4, after_pt=6, keep_next=True)
    _make('AI_Labels',       'Calibri',  9, after_pt=4)
    _make('AI_Blockquote',   'Georgia', 10, italic=True, before_pt=6, after_pt=6)
    _make('AI_SectionHeader','Calibri',  8, bold=True,  before_pt=10, after_pt=3)
    # Chapter opener
    _make('AI_ChapterTitle', 'Georgia', 22, bold=True,  before_pt=0, after_pt=12,
          page_break_before=True)
    _make('AI_ChapterIntro', 'Georgia', 11, italic=True, before_pt=0, after_pt=8)
    _make('AI_EntryListItem','Calibri', 10, after_pt=3)
    # Front/back matter
    _make('AI_HalfTitle',    'Georgia', 28, bold=True,  align=WD_ALIGN_PARAGRAPH.CENTER,
          before_pt=80, after_pt=8)
    _make('AI_SubTitle',     'Georgia', 13, italic=True,align=WD_ALIGN_PARAGRAPH.CENTER,
          after_pt=6)
    _make('AI_CopyrightBody','Georgia', 10, after_pt=5, line_spacing_pt=13)
    _make('AI_ForewordH1',   'Georgia', 18, bold=True,  after_pt=10)
    _make('AI_ForewordSH',   'Calibri',  8, bold=True,  before_pt=12, after_pt=4,
          color_rgb=(80,80,80))
    _make('AI_AbbrTerm',     'Georgia', 10, bold=True,  after_pt=2)
    _make('AI_AbbrDef',      'Georgia', 10, after_pt=5)
    _make('AI_IndexH1',      'Georgia', 18, bold=True,  after_pt=10)
    _make('AI_IndexLetter',  'Georgia', 12, bold=True,  before_pt=10, after_pt=2)
    _make('AI_IndexChapter', 'Georgia', 11, bold=True,  before_pt=4,  after_pt=1)
    _make('AI_IndexEntry',   'Georgia', 10, after_pt=1)
    _make('AI_PartLabel',    'Calibri',  8, color_rgb=(140,100,50), after_pt=4)
    _make('AI_PartTitle',    'Georgia', 24, bold=True,  after_pt=10)
    _make('AI_PartIntro',    'Georgia', 11, after_pt=8)
    _make('AI_SourceLabel',  'Calibri',  8, color_rgb=(140,100,50), after_pt=4)
    _make('AI_SourceTitle',  'Georgia', 18, bold=True,  after_pt=10)
    _make('AI_TOCHeading',   'Georgia', 18, bold=True,  after_pt=10)
    _make('AI_TOCEntry',     'Calibri', 11, after_pt=3)

    # Blockquote indentation (must be set after style creation)
    bq = doc.styles['AI_Blockquote']
    bq.paragraph_format.left_indent  = Mm(12)
    bq.paragraph_format.right_indent = Mm(12)


# ── Document structure helpers ─────────────────────────────────────────────────

def add_page_break(doc):
    """Add an explicit page break (blank page)."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def set_section_page_numbering(section, fmt='lowerRoman', start=1):
    """Set page number format and start for a section."""
    sectPr = section._sectPr
    for old in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(old)
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), fmt)
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)


def add_section_break_next_page(doc):
    """Insert a next-page section break and return the new section."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    pgType = OxmlElement('w:type')
    pgType.set(qn('w:val'), 'nextPage')
    sectPr.append(pgType)
    pPr.append(sectPr)
    return doc.sections[-1]


# ── Front matter ───────────────────────────────────────────────────────────────

def add_half_title(doc):
    """Page i — half-title."""
    doc.add_paragraph('Analyzing Islam', style='AI_HalfTitle')
    doc.add_paragraph('Volume I — The Quran', style='AI_SubTitle')
    doc.add_paragraph('A Critical Reference Guide', style='AI_SubTitle')


def add_copyright(doc):
    """Page iii — copyright."""
    doc.add_paragraph('Analyzing Islam — Volume I: The Quran', style='AI_CopyrightBody')
    doc.add_paragraph('A Critical Reference Guide', style='AI_CopyrightBody')
    doc.add_paragraph('', style='AI_CopyrightBody')
    doc.add_paragraph('© 2026 Analyzing Islam. All rights reserved.', style='AI_CopyrightBody')
    doc.add_paragraph('analyzingislam.com', style='AI_CopyrightBody')
    doc.add_paragraph('', style='AI_CopyrightBody')
    doc.add_paragraph('First edition, 2026.', style='AI_CopyrightBody')
    doc.add_paragraph('', style='AI_CopyrightBody')
    p = doc.add_paragraph(style='AI_CopyrightBody')
    p.add_run('All Quranic verses are quoted from the ')
    r = p.add_run('Saheeh International')
    r.italic = True
    p.add_run((' English translation — the Saudi-sanctioned mainstream Sunni edition, '
               'widely used in mosques and Islamic universities across the English-speaking world. '
               'This volume covers the Quran exclusively. Hadith collections — '
               'Sahih al-Bukhari, Sahih Muslim, and the four Sunan — are examined in subsequent volumes.'))
    doc.add_paragraph('', style='AI_CopyrightBody')
    doc.add_paragraph(('No part of this publication may be reproduced or transmitted in any form '
                        'without prior written permission from the publisher, except for brief quotations '
                        'in reviews or scholarly work with full attribution.'),
                       style='AI_CopyrightBody')
    doc.add_paragraph('', style='AI_CopyrightBody')
    doc.add_paragraph('Every entry references a specific verse — verify before citing.',
                       style='AI_CopyrightBody')
    doc.add_paragraph('', style='AI_CopyrightBody')
    doc.add_paragraph('ISBN — [to be assigned]', style='AI_CopyrightBody')


def add_toc(doc):
    """Pages v–vi — Table of Contents (Word auto-generates on open)."""
    doc.add_paragraph('Contents', style='AI_TOCHeading')

    # Insert Word TOC field — user must press F9 / Update Table on first open
    para = doc.add_paragraph(style='AI_TOCEntry')
    run = para.add_run()
    r = run._r

    fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin'); r.append(fc1)

    run2 = para.add_run()
    it = OxmlElement('w:instrText')
    it.set(qn('xml:space'), 'preserve')
    it.text = r' TOC \o "1-1" \h \z \u '
    run2._r.append(it)

    run3 = para.add_run()
    fc3 = OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'), 'separate'); run3._r.append(fc3)

    run4 = para.add_run('Right-click → Update Field after opening to generate page numbers.')

    run5 = para.add_run()
    fc5 = OxmlElement('w:fldChar'); fc5.set(qn('w:fldCharType'), 'end'); run5._r.append(fc5)


def add_foreword(doc):
    """Pages vii–ix — 3-page foreword."""
    doc.add_paragraph('Foreword', style='AI_ForewordH1')

    doc.add_paragraph('WHAT THIS IS', style='AI_ForewordSH')
    doc.add_paragraph(
        'This book is a reference catalog of passages from the Quran that present philosophical, '
        'historical, moral, or logical difficulties. Every entry cites a specific verse, explains '
        'what it says in plain language, and builds the case for why it presents a problem. Where '
        'relevant, entries also present the standard Muslim apologetic response and push back on it.',
        style='AI_Normal')
    doc.add_paragraph(
        'Volume I covers one source: the Quran — the text that Islam itself holds to be the direct, '
        'unaltered word of Allah, superior to all other Islamic sources in authority. No fringe '
        'interpretations. No hostile translations. The case is built entirely from Islam\'s own most '
        'authoritative scripture as rendered in its most widely endorsed English edition.',
        style='AI_Normal')

    doc.add_paragraph('HOW ENTRIES ARE ORGANISED', style='AI_ForewordSH')
    doc.add_paragraph(
        'Entries are grouped into 23 thematic chapters — Abrogation, Contradictions, Warfare & Jihad, '
        'Women & Sexual Issues, and so on. Each chapter collects every entry that belongs to that theme. '
        'If the Quran yields no entries in a given category, that chapter is omitted entirely.',
        style='AI_Normal')
    doc.add_paragraph(
        'When an entry touches more than one theme, it appears under the category that best captures '
        'its primary problem. It does not appear twice. Subsequent volumes examine the hadith '
        'collections — Sahih al-Bukhari, Sahih Muslim, and the four canonical Sunan.',
        style='AI_Normal')

    doc.add_paragraph('HOW TO READ AN ENTRY', style='AI_ForewordSH')
    doc.add_paragraph('Each entry contains four elements:', style='AI_Normal')
    for label, desc in [
        ('REFERENCE', 'The Quran citation — surah and verse number.'),
        ('RATING',    'The apologetic difficulty level: Basic, Moderate, or Strong.'),
        ('PASSAGE',   'The verse quoted in full from the Saheeh International translation.'),
        ('COMMENTARY','An explanation of the problem — what it says, why it matters, '
                       'and where the standard apologetic falls short.'),
    ]:
        p = doc.add_paragraph(style='AI_Normal')
        run = p.add_run(label + '  ')
        run.bold = True
        p.add_run(desc)

    # Page 2
    doc.add_paragraph('STRENGTH RATINGS', style='AI_ForewordSH')
    doc.add_paragraph(
        'Every entry is rated according to how difficult the problem is to answer from within the '
        'Islamic apologetic tradition:', style='AI_Normal')
    for level, desc in [
        ('Basic',    'Apologists have a stock reply. The problem is real but the standard response '
                     'is widely known and rehearsed.'),
        ('Moderate', 'Answering requires conceding something — softening a claim, reinterpreting a text, '
                     'or acknowledging that the tradition is not unanimous.'),
        ('Strong',   'The apologetic moves themselves generate new problems. Every standard response '
                     'either contradicts another Islamic claim or requires abandoning the plain meaning '
                     'of the text.'),
    ]:
        p = doc.add_paragraph(style='AI_Normal')
        run = p.add_run(level + '  ')
        run.bold = True
        p.add_run(desc)
    doc.add_paragraph(
        'Ratings reflect apologetic difficulty — not moral severity. A passage can be morally '
        'disturbing and still rated Basic if the apologetic reply is well-established and coherent.',
        style='AI_Normal')

    doc.add_paragraph('SOURCES AND TRANSLATIONS', style='AI_ForewordSH')
    doc.add_paragraph(
        'All Quranic verses in this volume are quoted from the Saheeh International English '
        'translation — the Saudi-sanctioned mainstream Sunni edition, widely used in mosques and '
        'Islamic universities across the English-speaking world. It is the edition most commonly '
        'recommended by contemporary Sunni scholars when asked to name an accurate English Quran.',
        style='AI_Normal')
    doc.add_paragraph(
        'Choosing the most mainstream, most-recommended translation removes the easy dismissal of '
        '"hostile translation." The problems documented in this volume are not artifacts of a '
        'tendentious rendering — they appear in the text that Islam\'s own authorities have endorsed '
        'and distributed worldwide.',
        style='AI_Normal')

    # Page 3
    doc.add_paragraph('A NOTE ON TONE', style='AI_ForewordSH')
    doc.add_paragraph(
        'This catalog does not argue. It presents. The entries speak through the texts themselves — '
        'the reader is left to draw their own conclusions. No passage is fabricated, paraphrased to '
        'distort, or stripped of context that would change its meaning. Where context matters, it is '
        'provided.',
        style='AI_Normal')
    doc.add_paragraph(
        'The commentary aims to be precise rather than polemical. Where Islamic scholars disagree '
        'among themselves, that disagreement is noted. Where a passage has a defensible reading, '
        'that reading is acknowledged before the problem with it is explained. The goal is not to '
        'mock but to examine — carefully, specifically, and without concession.',
        style='AI_Normal')
    doc.add_paragraph(
        'Readers who find a specific entry inaccurate, mistranslated, or missing essential context '
        'are encouraged to raise the objection at analyzingislam.com, where every entry in this '
        'volume is also published online and open to scrutiny.',
        style='AI_Normal')

    doc.add_paragraph('HOW TO USE THIS BOOK', style='AI_ForewordSH')
    doc.add_paragraph(
        'Read it in order or jump directly to a category. Each entry stands on its own. The 23 '
        'thematic chapters are self-contained — no prior entry is assumed when reading any later entry.',
        style='AI_Normal')
    doc.add_paragraph(
        'Use the Quran Verse Index at the back to locate entries by surah and verse number. Use the '
        'General Index to find entries touching a specific topic, person, or concept across all 23 '
        'chapters. Every entry references a specific verse. Verify before citing.',
        style='AI_Normal')


def add_abbreviations(doc):
    """Pages x–xi — Abbreviations & Reference Guide."""
    doc.add_paragraph('Abbreviations & Reference Guide', style='AI_ForewordH1')

    doc.add_paragraph('CITATION FORMAT', style='AI_ForewordSH')
    p = doc.add_paragraph(style='AI_Normal')
    run = p.add_run('Q 4:34  ')
    run.bold = True
    p.add_run('Quran, Surah 4 (An-Nisa), Verse 34. All Quranic citations follow this '
              'surah:verse format. Where a range of verses is relevant, it appears as Q 9:5–6. '
              'All quotations are from the Saheeh International English translation.')

    doc.add_paragraph('STRENGTH RATINGS', style='AI_ForewordSH')
    for term, defn in [
        ('Basic',    'Apologists have a stock reply. The problem is real but the standard response '
                     'is widely known and rehearsed.'),
        ('Moderate', 'Answering requires conceding something — softening a claim or reinterpreting '
                     'the text.'),
        ('Strong',   'Apologetic moves generate new problems. Every standard response requires '
                     'abandoning the plain meaning of the text or contradicts another Islamic claim.'),
    ]:
        p = doc.add_paragraph(style='AI_Normal')
        run = p.add_run(term + '  ')
        run.bold = True
        p.add_run(defn)

    doc.add_paragraph('QURANIC TERMINOLOGY', style='AI_ForewordSH')
    quranic_terms = [
        ('Ayah (pl. Ayat)', 'A verse of the Quran; literally "a sign"'),
        ('Surah', 'A chapter of the Quran; there are 114 in total'),
        ('Meccan', 'Revealed while Muhammad was in Mecca (c. 610–622 CE) — generally monotheism and eschatology'),
        ('Medinan', 'Revealed while Muhammad was in Medina (c. 622–632 CE) — generally law, governance, and warfare'),
        ('Naskh', 'Abrogation — the doctrine that later Quranic verses can cancel earlier ones'),
        ('Tafsir', 'Quranic exegesis or commentary; the classical tradition of explaining individual verses'),
        ('Asbab al-Nuzul', 'The "occasions of revelation" — historical circumstances that triggered specific verses'),
    ]
    for term, defn in quranic_terms:
        p = doc.add_paragraph(style='AI_Normal')
        run = p.add_run(term + '  ')
        run.bold = True
        p.add_run(defn)

    # Page 2
    doc.add_paragraph('ARABIC & ISLAMIC TERMINOLOGY', style='AI_ForewordSH')
    arabic_terms = [
        ('Fiqh', 'Islamic jurisprudence — the body of legal rulings derived from the Quran and hadith'),
        ('Ulema', 'Islamic scholars and jurists collectively'),
        ('Dhimmi', 'A non-Muslim subject living under Islamic rule, subject to the jizya tax and legal restrictions'),
        ('Hudud', 'Fixed Quranic punishments — amputation, stoning, lashing — that cannot be reduced by a judge'),
        ('Jizya', 'A tax levied on non-Muslims living under Islamic governance in lieu of military service (Q 9:29)'),
        ('Jinn', 'Supernatural beings made of smokeless fire; mentioned throughout the Quran'),
        ('Tahrif', 'The Islamic claim that Jews and Christians corrupted their scriptures'),
        ('Fitnah', 'Trial, strife, or persecution; used in key warfare verses (Q 2:193, 8:39)'),
        ('Ma malakat aymanukum', '"What your right hands possess" — the Quranic phrase for enslaved people and captive women'),
        ('Dahaha', 'Verb in Q 79:30 translated "spread out" or "egg-shaped" — key in cosmology debates'),
        ('Makr', 'Plotting or scheming; used of Allah in Q 3:54 and 8:30'),
        ('Ijaz al-Quran', 'The doctrine of the Quran\'s inimitability — the claim that its literary style is miraculous'),
        ('r.a.', 'Radi Allahu anhu / anha — "May Allah be pleased with him / her"'),
        ('s.a.w.', 'Sallallahu alayhi wa sallam — "Peace and blessings be upon him"'),
    ]
    for term, defn in arabic_terms:
        p = doc.add_paragraph(style='AI_Normal')
        run = p.add_run(term + '  ')
        run.bold = True
        p.add_run(defn)


# ── Body content ───────────────────────────────────────────────────────────────

def add_part_opener(doc):
    """The Quran part opener — first page of body (arabic p.1)."""
    doc.add_paragraph('VOLUME I', style='AI_PartLabel')
    doc.add_paragraph('The Quran', style='AI_PartTitle')
    doc.add_paragraph(
        'Comprising 114 surahs and 6,236 verses, the Quran is Islam\'s central divine text — '
        'believed by Muslims to be the literal word of Allah as revealed to Muhammad between '
        'approximately 610 and 632 CE. The surahs were revealed in Mecca and Medina, collected '
        'under Abu Bakr and standardised under Uthman ibn Affan around 650 CE. All verses in '
        'this volume are quoted from the Saheeh International English translation.',
        style='AI_PartIntro')
    p = doc.add_paragraph(style='AI_Normal')
    run = p.add_run('262 ENTRIES ACROSS 23 CHAPTERS')
    run.bold = True


def add_source_intro(doc):
    """Source introduction page."""
    add_page_break(doc)
    doc.add_paragraph('PRIMARY SOURCE', style='AI_SourceLabel')
    doc.add_paragraph('The Quran', style='AI_SourceTitle')
    doc.add_paragraph(
        'The Quran is the central religious text of Islam, believed by Muslims to be the direct '
        'word of Allah as revealed to the Prophet Muhammad through the angel Jibril (Gabriel) over '
        'approximately twenty-three years — from 610 CE until Muhammad\'s death in 632 CE. It '
        'comprises 114 chapters (surahs) containing 6,236 verses (ayahs), arranged roughly in '
        'descending order of length rather than chronological order of revelation.',
        style='AI_Normal')
    doc.add_paragraph(
        'During Muhammad\'s lifetime, verses were memorised by companions and recorded on various '
        'materials. Following the Battle of Yamama in 632 CE, in which many memorisers were killed, '
        'Caliph Abu Bakr commissioned a written compilation. The standardised text known today was '
        'established under the third Caliph Uthman ibn Affan around 650 CE; variant readings were '
        'officially destroyed.',
        style='AI_Normal')
    doc.add_paragraph(
        'Surahs revealed in Mecca — primarily the earlier, shorter chapters — tend to focus on '
        'monotheism, eschatology, and moral instruction. Those revealed in Medina — longer, later '
        'chapters — deal extensively with law, governance, warfare, and relations with non-Muslims. '
        'The doctrine of abrogation (naskh) holds that later verses may supersede earlier ones, a '
        'principle with significant ethical implications examined throughout this volume.',
        style='AI_Normal')
    doc.add_paragraph(
        'All Quranic verses in this volume are quoted from the Saheeh International English '
        'translation, the edition most widely recommended by contemporary Sunni scholars for '
        'accuracy to the Arabic.',
        style='AI_Normal')
    # Stats row
    p = doc.add_paragraph(style='AI_Normal')
    for stat, label in [('114', 'SURAHS'), ('6,236', 'VERSES'),
                        ('610–632 CE', 'REVELATION PERIOD'), ('c. 650 CE', "UTHMAN'S COMPILATION")]:
        run = p.add_run(f'{stat}  ')
        run.bold = True
        p.add_run(f'{label}    ')


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    entries  = get_entries()
    content  = parse_entries()
    chapters = build_chapters(entries)

    print(f"Entries: {len(entries)}  |  Chapters: {sum(1 for c in chapters.values() if c)}")

    doc = Document()
    setup_document(doc)
    setup_styles(doc)

    # ── Front matter (roman numerals: i, ii, iii…) ──────────────────────────
    set_section_page_numbering(doc.sections[0], fmt='lowerRoman', start=1)

    add_half_title(doc)    # page i
    add_page_break(doc)    # page ii (blank)
    add_copyright(doc)     # page iii
    add_page_break(doc)    # page iv (blank)

    add_toc(doc)           # pages v–vi
    add_foreword(doc)      # pages vii–ix
    add_abbreviations(doc) # pages x–xi

    # ── Body section (arabic numerals: 1, 2, 3…) ────────────────────────────
    body_section = add_section_break_next_page(doc)
    setup_footer(body_section)
    set_section_page_numbering(body_section, fmt='decimal', start=1)

    add_part_opener(doc)
    add_source_intro(doc)

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
